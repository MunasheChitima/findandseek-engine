"""DOCX paragraphs, tables + embedded images."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from find_and_seek.ingest.extract.base import ExtractResult, ImageRegion, TextBlock


def extract_docx(path: Path) -> ExtractResult:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    blocks: list[TextBlock] = []
    if paragraphs:
        blocks.append(
            TextBlock(
                text="\n\n".join(paragraphs),
                source_type="text",
                location_ref="document",
            )
        )

    # Tables are appended after the paragraphs block — python-docx can't cheaply
    # interleave true document order.
    for i, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                text = cell.text.strip()
                # Merged cells repeat the same cell object; collapse consecutive dupes.
                if cells and cells[-1] == text:
                    continue
                cells.append(text)
            line = " | ".join(cells).strip()
            if line.strip(" |"):
                rows.append(line)
        if rows:
            blocks.append(
                TextBlock(
                    text="\n".join(rows),
                    source_type="table",
                    location_ref=f"table {i + 1}",
                )
            )

    images: list[ImageRegion] = []
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            try:
                blob = rel.target_part.blob
                images.append(
                    ImageRegion(
                        data=blob,
                        mime="image/png",
                        location_ref="embedded image",
                    )
                )
            except (AttributeError, KeyError):
                continue

    return ExtractResult(blocks=blocks, images=images, file_type="document")
