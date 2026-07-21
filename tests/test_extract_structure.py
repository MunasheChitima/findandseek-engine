"""Structural extraction: DOCX tables + magic-bytes format sniffing."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from find_and_seek.ingest.extract.router import extract_file, sniff_format


def _make_docx(path: Path, *, merged: bool = False) -> None:
    doc = Document()
    doc.add_paragraph("Intro paragraph before the table.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Plumbing"
    table.cell(1, 1).text = "450.00"
    if merged:
        merged_cell = table.cell(0, 0).merge(table.cell(0, 1))
        merged_cell.text = "X"
    doc.save(str(path))


class TestDocxTables:
    def test_table_block_extracted(self, tmp_path):
        p = tmp_path / "invoice.docx"
        _make_docx(p)
        result = extract_file(p)
        tables = [b for b in result.blocks if b.source_type == "table"]
        assert tables, "expected a table block"
        assert "Item | Amount" in tables[0].text
        assert "Plumbing | 450.00" in tables[0].text
        assert tables[0].location_ref == "table 1"

    def test_paragraphs_still_extracted(self, tmp_path):
        p = tmp_path / "invoice.docx"
        _make_docx(p)
        result = extract_file(p)
        text_blocks = [b for b in result.blocks if b.source_type == "text"]
        assert text_blocks and "Intro paragraph" in text_blocks[0].text

    def test_merged_cells_collapsed(self, tmp_path):
        p = tmp_path / "merged.docx"
        _make_docx(p, merged=True)
        result = extract_file(p)
        tables = [b for b in result.blocks if b.source_type == "table"]
        assert tables
        assert "X | X" not in tables[0].text
        assert "X" in tables[0].text


class TestSniffFormat:
    def test_pdf(self, tmp_path):
        p = tmp_path / "doc.bin"
        p.write_bytes(b"%PDF-1.4 fake content")
        assert sniff_format(p) == ".pdf"

    def test_rtf(self, tmp_path):
        p = tmp_path / "doc.bin"
        p.write_bytes(b"{\\rtf1\\ansi Hello}")
        assert sniff_format(p) == ".rtf"

    def test_docx_zip(self, tmp_path):
        p = tmp_path / "doc.bin"
        with zipfile.ZipFile(p, "w") as zf:
            zf.writestr("word/document.xml", "<w:document/>")
        assert sniff_format(p) == ".docx"


class TestMislabeledDispatch:
    def test_rtf_named_docx(self, tmp_path):
        p = tmp_path / "wrong.docx"
        p.write_bytes(b"{\\rtf1\\ansi Hello mislabeled world}")
        result = extract_file(p)
        text = "\n".join(b.text for b in result.blocks)
        assert "Hello mislabeled world" in text

    def test_docx_named_txt(self, tmp_path):
        # A real docx saved under a .txt suffix used to skip sniffing entirely
        # (.txt is a "recognized" extension) and get force-decoded as text via
        # errors="replace" — the ZIP bytes "succeed" as tens of thousands of
        # mangled garbage characters instead of routing to extract_docx.
        p = tmp_path / "wrong.txt"
        _make_docx(p)
        result = extract_file(p)
        tables = [b for b in result.blocks if b.source_type == "table"]
        assert tables, "expected the real docx table, not garbage-decoded zip bytes"
        assert "Item | Amount" in tables[0].text
        assert "Plumbing | 450.00" in tables[0].text

    def test_docx_named_csv(self, tmp_path):
        p = tmp_path / "wrong.csv"
        _make_docx(p)
        result = extract_file(p)
        text = "\n".join(b.text for b in result.blocks)
        assert "PK\x03\x04" not in text
        assert "Plumbing" in text

    def test_xlsx_named_txt(self, tmp_path):
        # openpyxl's load_workbook validates the FILENAME suffix against a
        # fixed allowlist before touching content, independent of router.py's
        # content-based sniff — a real xlsx saved under .txt raised "openpyxl
        # does not support .txt file format" even after routing correctly.
        from openpyxl import Workbook

        p = tmp_path / "wrong.txt"
        wb = Workbook()
        ws = wb.active
        ws.append(["Item", "Amount"])
        ws.append(["Plumbing", 450.00])
        wb.save(str(p))

        result = extract_file(p)
        text = "\n".join(b.text for b in result.blocks)
        assert "Plumbing" in text
        assert "450" in text

    def test_jpeg_named_txt(self, tmp_path):
        # A JFIF/JPEG saved under a .txt suffix used to sail through
        # extract_text_file's errors="replace" decode as garbage control-char
        # text instead of being recognized as an image.
        p = tmp_path / "wrong.txt"
        p.write_bytes(b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00" + b"\x00" * 64)
        assert sniff_format(p) == ".jpg"
        result = extract_file(p)
        assert result.file_type == "image"
        assert result.images and result.images[0].mime == "image/jpeg"


class TestEmlHeaderFormatting:
    def test_from_to_render_as_readable_addresses(self, tmp_path):
        # mailparser returns from/to as a list of (display_name, address)
        # tuples — a naive f-string over that leaks the raw Python repr
        # (e.g. "From: [('', 'sender@example.com')]") into extracted text,
        # which then gets fed straight into summarisation.
        p = tmp_path / "msg.eml"
        p.write_text(
            "From: Jane Doe <jane@example.com>\n"
            "To: John Smith <john@example.com>\n"
            "Subject: Transaction Details\n"
            "Date: Sun, 15 Mar 2026 10:00:00 +0000\n"
            "Content-Type: text/plain\n"
            "\n"
            "Please review the attached details.\n"
        )
        from find_and_seek.ingest.extract.email import extract_eml

        result = extract_eml(p)
        text = result.blocks[0].text
        assert "[('" not in text
        assert "Jane Doe <jane@example.com>" in text
        assert "John Smith <john@example.com>" in text
